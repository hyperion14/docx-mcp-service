"""BHK Format Converter for DOCX Generation.

Converts Markdown text to BHK-formatted Word documents without Markdown syntax.
Based on paragraph_converter.py logic from bhk-md-converter.
"""
import logging
import re
from typing import Optional, Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)

try:
    import mistune
    HAS_MISTUNE = True
except ImportError:
    HAS_MISTUNE = False
    logger.warning("mistune not installed - BHK formatting will be limited")


class BHKFormatter:
    """Converts Markdown to BHK-formatted DOCX.
    
    Features (based on paragraph_converter.py Mode B):
    - Headings → BHK_Standard paragraphs (bold, no auto-numbering)
    - List items → BHK_Standard paragraphs with bullet markers (•)
    - Manual numbering preserved (e.g., "## 1. Title" stays "1. Title")
    - Flat structure (no hierarchical numbering)
    - Inline formatting preserved (**bold**, _italic_)
    """
    
    def __init__(self, template_path: Optional[str] = None):
        """Initialize BHK Formatter.
        
        Args:
            template_path: Path to BHK template DOCX (must contain BHK_Standard style)
        """
        self.template_path = template_path
        
    def convert_to_docx(self, doc: Document, markdown_text: str) -> None:
        """Convert Markdown text to BHK-formatted DOCX.
        
        Args:
            doc: Document object to populate
            markdown_text: Markdown-formatted text
        """
        if not HAS_MISTUNE:
            # Fallback: Simple line-based conversion
            logger.warning("Using fallback converter (mistune not available)")
            self._convert_simple(doc, markdown_text)
            return
            
        try:
            # Parse Markdown to AST
            md = mistune.create_markdown(renderer='ast')
            ast = md(markdown_text)
            
            # Process AST nodes
            self._process_nodes(doc, ast)
            
        except Exception as e:
            logger.error(f"Markdown parsing failed: {e}", exc_info=True)
            # Fallback to simple conversion
            self._convert_simple(doc, markdown_text)
    
    def _process_nodes(self, doc: Document, nodes: List[Dict[str, Any]]) -> None:
        """Process AST nodes and add to document.
        
        Args:
            doc: Document object
            nodes: List of AST nodes from mistune
        """
        for node in nodes:
            node_type = node.get('type', '')
            
            if node_type == 'heading':
                self._process_heading(doc, node)
            elif node_type == 'list':
                self._process_list(doc, node)
            elif node_type == 'paragraph':
                self._process_paragraph(doc, node)
            elif node_type == 'thematic_break':
                # Ignore --- separators (common in Markdown)
                pass
            elif node_type == 'block_code':
                # Code blocks as plain BHK_Standard paragraphs
                self._process_code_block(doc, node)
            else:
                logger.debug(f"Unhandled node type: {node_type}")
    
    def _process_heading(self, doc: Document, node: Dict[str, Any]) -> None:
        """Process heading as BHK_Standard paragraph.
        
        Mode B: Keeps manual numbering and uses BHK_Standard style with
        bold formatting to distinguish headings.
        
        Args:
            doc: Document object
            node: Heading node from AST
        """
        level = node.get('attrs', {}).get('level', 1)
        children = node.get('children', [])
        
        logger.debug(f"Processing H{level} as BHK_Standard (bold)")
        
        # Create paragraph
        para = doc.add_paragraph()
        
        # Apply formatted text (preserves existing bold/underline)
        self._apply_inline_formatting(para, children)
        
        # Make entire heading bold to distinguish from regular paragraphs
        for run in para.runs:
            if not run.bold:
                run.bold = True
        
        # Apply BHK_Standard style
        self._apply_bhk_style(para)
        
        # Optional: Add spacing after headings
        para.paragraph_format.space_after = Pt(6)
    
    def _process_list(self, doc: Document, node: Dict[str, Any], depth: int = 0) -> None:
        """Process list as BHK_Standard paragraphs with bullet markers.
        
        Mode B: Uses BHK_Standard style with manual bullet markers (•)
        instead of auto-numbered lists.
        
        Args:
            doc: Document object
            node: List node from AST
            depth: Nesting depth (0 for top-level)
        """
        children = node.get('children', [])
        is_ordered = node.get('attrs', {}).get('ordered', False)
        
        for idx, item in enumerate(children, 1):
            if item.get('type') != 'list_item':
                continue
                
            item_children = item.get('children', [])
            
            # Create paragraph
            para = doc.add_paragraph()
            
            # Add bullet marker or number
            if is_ordered:
                para.add_run(f'{idx}. ')
            else:
                para.add_run('• ')
            
            # Process item content (may contain nested lists)
            for child in item_children:
                if child.get('type') == 'list':
                    # Nested list - process recursively
                    self._process_list(doc, child, depth + 1)
                elif child.get('type') == 'paragraph':
                    # Inline content
                    self._apply_inline_formatting(para, child.get('children', []))
                else:
                    # Other content
                    self._apply_inline_formatting(para, [child])
            
            # Apply BHK_Standard style
            self._apply_bhk_style(para)
            
            # Set indentation for nested items
            if depth > 0:
                para.paragraph_format.left_indent = Inches(0.25 * depth)
    
    def _process_paragraph(self, doc: Document, node: Dict[str, Any]) -> None:
        """Process standard paragraph as BHK_Standard.
        
        Args:
            doc: Document object
            node: Paragraph node from AST
        """
        children = node.get('children', [])
        
        # Skip empty paragraphs
        text_content = self._extract_text(children)
        if not text_content.strip():
            return
        
        # Create paragraph
        para = doc.add_paragraph()
        
        # Apply formatted text
        self._apply_inline_formatting(para, children)
        
        # Apply BHK_Standard style
        self._apply_bhk_style(para)
    
    def _process_code_block(self, doc: Document, node: Dict[str, Any]) -> None:
        """Process code block as plain BHK_Standard paragraph.
        
        Args:
            doc: Document object
            node: Code block node from AST
        """
        raw_code = node.get('raw', '')
        
        para = doc.add_paragraph(raw_code)
        self._apply_bhk_style(para)
        
        # Optional: Use monospace font for code
        for run in para.runs:
            run.font.name = 'Courier New'
    
    def _apply_inline_formatting(self, para, children: List[Dict[str, Any]]) -> None:
        """Apply inline formatting (bold, italic, text) to paragraph.
        
        Args:
            para: Paragraph object
            children: List of inline nodes (text, strong, emphasis, etc.)
        """
        for child in children:
            child_type = child.get('type', '')
            
            if child_type == 'text':
                # Plain text
                raw_text = child.get('raw', '')
                para.add_run(raw_text)
                
            elif child_type == 'strong':
                # **bold**
                text = self._extract_text(child.get('children', []))
                run = para.add_run(text)
                run.bold = True
                
            elif child_type == 'emphasis':
                # _italic_
                text = self._extract_text(child.get('children', []))
                run = para.add_run(text)
                run.italic = True
                
            elif child_type == 'codespan':
                # `code`
                code_text = child.get('raw', '')
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                
            elif child_type == 'link':
                # [text](url)
                link_text = self._extract_text(child.get('children', []))
                url = child.get('attrs', {}).get('url', '')
                run = para.add_run(f'{link_text} ({url})')
                run.font.underline = True
                
            elif child_type == 'linebreak':
                # Line break
                para.add_run('\n')
                
            else:
                # Unknown inline type - extract text
                text = self._extract_text([child])
                if text:
                    para.add_run(text)
    
    def _extract_text(self, children: List[Dict[str, Any]]) -> str:
        """Extract plain text from AST nodes.
        
        Args:
            children: List of AST nodes
            
        Returns:
            Plain text string
        """
        text_parts = []
        
        for child in children:
            if isinstance(child, str):
                text_parts.append(child)
            elif child.get('type') == 'text':
                text_parts.append(child.get('raw', ''))
            elif 'children' in child:
                # Recursive extraction
                text_parts.append(self._extract_text(child['children']))
            elif 'raw' in child:
                text_parts.append(child['raw'])
                
        return ''.join(text_parts)
    
    def _apply_bhk_style(self, para) -> None:
        """Apply BHK_Standard style to paragraph.
        
        Falls back to Normal if BHK_Standard not available.
        
        Args:
            para: Paragraph object
        """
        try:
            # Try to apply BHK_Standard style
            para.style = 'BHK_Standard'
            logger.debug("Applied BHK_Standard style")
        except KeyError:
            # Fallback: Use Normal style
            logger.warning("BHK_Standard style not found in template, using Normal")
            try:
                para.style = 'Normal'
            except KeyError:
                logger.warning("Normal style also not found, using default")
    
    def _convert_simple(self, doc: Document, text: str) -> None:
        """Fallback: Simple line-based conversion without full Markdown parsing.
        
        Uses regex to detect headings, lists, bold, italic.
        
        Args:
            doc: Document object
            text: Plain or Markdown text
        """
        lines = text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Skip separator lines
            if line.strip() in ['---', '***', '___']:
                i += 1
                continue
            
            # Detect headings (###, ##, #)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2)
                
                para = doc.add_paragraph()
                self._apply_simple_inline_formatting(para, heading_text)
                
                # Make bold
                for run in para.runs:
                    run.bold = True
                
                self._apply_bhk_style(para)
                para.paragraph_format.space_after = Pt(6)
                i += 1
                continue
            
            # Detect list items (-, *, +)
            list_match = re.match(r'^[\s]*[-*+]\s+(.+)$', line)
            if list_match:
                list_text = list_match.group(1)
                
                para = doc.add_paragraph()
                para.add_run('• ')
                self._apply_simple_inline_formatting(para, list_text)
                self._apply_bhk_style(para)
                i += 1
                continue
            
            # Detect numbered lists (1., 2., etc.)
            numbered_match = re.match(r'^[\s]*(\d+)\.\s+(.+)$', line)
            if numbered_match:
                number = numbered_match.group(1)
                list_text = numbered_match.group(2)
                
                para = doc.add_paragraph()
                para.add_run(f'{number}. ')
                self._apply_simple_inline_formatting(para, list_text)
                self._apply_bhk_style(para)
                i += 1
                continue
            
            # Regular paragraph
            if line.strip():
                para = doc.add_paragraph()
                self._apply_simple_inline_formatting(para, line)
                self._apply_bhk_style(para)
            
            i += 1
    
    def _apply_simple_inline_formatting(self, para, text: str) -> None:
        """Apply simple inline formatting using regex.
        
        Supports **bold** and _italic_.
        
        Args:
            para: Paragraph object
            text: Text with Markdown inline formatting
        """
        # Simple regex-based processing
        # Pattern: **bold** or __bold__
        bold_pattern = r'\*\*(.+?)\*\*|__(.+?)__'
        # Pattern: *italic* or _italic_
        italic_pattern = r'\*(.+?)\*|_(.+?)_'
        
        # For simplicity, process in order and handle text/bold/italic
        # This is a simplified version - full implementation would be more complex
        
        remaining = text
        
        while remaining:
            # Check for bold
            bold_match = re.search(bold_pattern, remaining)
            italic_match = re.search(italic_pattern, remaining)
            
            # Find first match
            first_match = None
            match_type = None
            
            if bold_match and (not italic_match or bold_match.start() < italic_match.start()):
                first_match = bold_match
                match_type = 'bold'
            elif italic_match:
                first_match = italic_match
                match_type = 'italic'
            
            if first_match:
                # Add text before match
                before_text = remaining[:first_match.start()]
                if before_text:
                    para.add_run(before_text)
                
                # Add formatted text
                formatted_text = first_match.group(1) or first_match.group(2)
                run = para.add_run(formatted_text)
                
                if match_type == 'bold':
                    run.bold = True
                elif match_type == 'italic':
                    run.italic = True
                
                # Continue with remaining text
                remaining = remaining[first_match.end():]
            else:
                # No more formatting - add remaining text
                para.add_run(remaining)
                break
