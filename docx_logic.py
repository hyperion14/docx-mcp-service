import os
import uuid
import threading
import time
import shutil
import logging
from datetime import datetime
from docx import Document

# Import BHK formatter
try:
    from bhk_formatter import BHKFormatter
    HAS_BHK_FORMATTER = True
except ImportError:
    logger.warning("BHKFormatter not available - falling back to plain text mode")
    HAS_BHK_FORMATTER = False

# Configure logging
logger = logging.getLogger(__name__)

def generate_docx_from_text(text, output_folder, template_path=None, custom_filename=None, use_bhk_format=True):
    """
    Generate DOCX from text input.
    
    Args:
        text: Plain or Markdown text to convert
        output_folder: Directory to save output files
        template_path: Path to BHK template (optional)
        custom_filename: Custom filename (optional)
        use_bhk_format: If True, parse Markdown and apply BHK formatting (default: True)
    
    Returns:
        Tuple of (docx_filename, txt_filename)
    """
    try:
        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%y%m%d_%H%M")

        if custom_filename:
            # Clean custom filename
            custom_filename = custom_filename.replace('.docx', '').replace('.doc', '')
            # Remove special characters
            custom_filename = ''.join(c for c in custom_filename if c.isalnum() or c in (' ', '-', '_'))
            custom_filename = custom_filename.replace(' ', '_').strip('_')
            filename = f"{timestamp}_{custom_filename}.docx"
        else:
            # Try to extract a meaningful name from the text content
            lines = text.split('\n')
            first_line = ''
            
            # Find first non-empty line
            for line in lines:
                if line.strip():
                    first_line = line.strip()
                    break
            
            # Create filename from first line (max 30 chars)
            if first_line:
                # Remove ALL common markdown symbols
                cleaned = first_line
                # Remove markdown symbols: #, *, _, `, >, -, etc.
                for symbol in ['#', '*', '_', '`', '>', '-', '|', '[', ']', '(', ')']:
                    cleaned = cleaned.replace(symbol, '')
                cleaned = cleaned.strip()
                
                # Keep only alphanumeric, spaces, hyphens, underscores
                cleaned = ''.join(c for c in cleaned if c.isalnum() or c in (' ', '-', '_'))
                cleaned = cleaned.replace(' ', '_')
                
                # Limit length
                if len(cleaned) > 30:
                    cleaned = cleaned[:30]
                
                if cleaned:
                    filename = f"{timestamp}_{cleaned}.docx"
                else:
                    filename = f"{timestamp}_dokument.docx"
            else:
                filename = f"{timestamp}_dokument.docx"

        # Save source text for archival
        txt_filename = filename.replace('.docx', '.txt')
        txt_path = os.path.join(output_folder, txt_filename)

        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text)

        logger.info(f"Saved source text: {txt_filename}")

        # Generate DOCX document
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()

        # Choose formatting mode
        if use_bhk_format and HAS_BHK_FORMATTER:
            logger.info("Using BHK format conversion (Markdown-aware)")
            formatter = BHKFormatter(template_path)
            formatter.convert_to_docx(doc, text)
        else:
            logger.info("Using legacy plain text conversion")
            _legacy_text_conversion(doc, text)

        # Save DOCX file
        docx_path = os.path.join(output_folder, filename)
        doc.save(docx_path)

        logger.info(f"Generated DOCX: {filename}")

        return filename, txt_filename

    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}", exc_info=True)
        raise e


def _legacy_text_conversion(doc, text):
    """Legacy plain text conversion (backward compatibility).
    
    Args:
        doc: Document object to populate
        text: Plain text content
    """
    # Add title if available (extract from first line if it looks like a title)
    lines = text.split('\n')
    if lines and len(lines[0]) < 100 and not lines[0].startswith(' '):
        doc.add_heading(lines[0], level=1)
        content_text = '\n'.join(lines[1:])
    else:
        content_text = text

    # Add paragraphs (split by double newline for better formatting)
    paragraphs = content_text.split('\n\n')
    
    # Check if we can use the last paragraph if it's empty (common in templates)
    target_style = None
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        last_p = doc.paragraphs[-1]
        target_style = last_p.style
    else:
        last_p = None

    first_para_written = False
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
            
        if not first_para_written and last_p:
            # Use the existing empty paragraph
            last_p.add_run(para_text.strip())
            first_para_written = True
        else:
            # Add new paragraph, inheriting style if available
            doc.add_paragraph(para_text.strip(), style=target_style)

def archive_files_after_delay(docx_filename, txt_filename, upload_folder, archive_folder, delay=86400):
    """
    Archive source and generated files after delay period.
    """
    time.sleep(delay)

    try:
        # Create archive subdirectory for today
        archive_date = datetime.now().strftime("%y%m%d")
        daily_archive = os.path.join(archive_folder, archive_date)
        os.makedirs(daily_archive, exist_ok=True)

        # Move DOCX file
        docx_src = os.path.join(upload_folder, docx_filename)
        docx_dst = os.path.join(daily_archive, docx_filename)

        if os.path.exists(docx_src):
            shutil.move(docx_src, docx_dst)
            logger.info(f"Archived DOCX: {docx_filename} -> {daily_archive}")

        # Move TXT file
        txt_src = os.path.join(upload_folder, txt_filename)
        txt_dst = os.path.join(daily_archive, txt_filename)

        if os.path.exists(txt_src):
            shutil.move(txt_src, txt_dst)
            logger.info(f"Archived TXT: {txt_filename} -> {daily_archive}")

    except Exception as e:
        logger.error(f"Error archiving files: {str(e)}", exc_info=True)
